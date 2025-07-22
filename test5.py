# RAG Strategy and Token Analysis System
import os
import re
import openai
import numpy as np
import tiktoken
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
import matplotlib.pyplot as plt
from collections import defaultdict
import time

print("🚀 RAG Strategy and Token Analysis System Starting...")

# === CONFIG === #
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY") or input("Enter your OpenAI API key: ").strip()
MODEL_NAME = "gpt-4"

# Model token limits
MODEL_LIMITS = {
    "gpt-4": 8192,
    "gpt-4-32k": 32768,
    "gpt-4-turbo": 128000,
    "gpt-3.5-turbo": 4096,
    "gpt-3.5-turbo-16k": 16384
}

# === TOKEN ANALYSIS FUNCTIONS === #
def count_tokens(text, model="gpt-4"):
    """Count tokens for specific model"""
    try:
        if model.startswith("gpt-4"):
            encoding = tiktoken.encoding_for_model("gpt-4")
        elif model.startswith("gpt-3.5"):
            encoding = tiktoken.encoding_for_model("gpt-3.5-turbo")
        else:
            encoding = tiktoken.get_encoding("cl100k_base")
        
        return len(encoding.encode(text))
    except:
        # Fallback to word approximation (1 token ≈ 0.75 words)
        return int(len(text.split()) * 1.33)

def analyze_token_distribution(chunks, model="gpt-4"):
    """Analyze token distribution across chunks"""
    token_counts = [count_tokens(chunk, model) for chunk in chunks]
    return {
        'total_tokens': sum(token_counts),
        'avg_tokens_per_chunk': np.mean(token_counts),
        'max_tokens_per_chunk': max(token_counts) if token_counts else 0,
        'min_tokens_per_chunk': min(token_counts) if token_counts else 0,
        'token_distribution': token_counts
    }

# === ENHANCED TEXT EXTRACTION === #
def extract_text_robust(filepath):
    """Robust text extraction with detailed reporting"""
    print(f"📄 Extracting text from {filepath}...")
    
    if not os.path.exists(filepath):
        print(f"❌ File not found: {filepath}")
        return "", {}
    
    extraction_stats = {
        'method_used': None,
        'paragraphs_found': 0,
        'tables_found': 0,
        'raw_length': 0,
        'cleaned_length': 0,
        'extraction_time': 0
    }
    
    start_time = time.time()
    
    # Try python-docx first
    try:
        import docx
        doc = docx.Document(filepath)
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
                extraction_stats['paragraphs_found'] += 1
        
        # Extract tables
        for table in doc.tables:
            extraction_stats['tables_found'] += 1
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        content = '\n'.join(text_parts)
        extraction_stats['method_used'] = 'python-docx'
        extraction_stats['raw_length'] = len(content)
        
    except Exception as e:
        print(f"❌ python-docx failed: {e}")
        
        # Try mammoth fallback
        try:
            import mammoth
            with open(filepath, "rb") as f:
                result = mammoth.extract_raw_text(f)
            content = result.value
            extraction_stats['method_used'] = 'mammoth'
            extraction_stats['raw_length'] = len(content)
        except Exception as e2:
            print(f"❌ mammoth failed: {e2}")
            return "", extraction_stats
    
    # Clean text
    cleaned_content = re.sub(r'\s+', ' ', content).strip()
    extraction_stats['cleaned_length'] = len(cleaned_content)
    extraction_stats['extraction_time'] = time.time() - start_time
    
    print(f"✅ Extracted {len(cleaned_content.split())} words using {extraction_stats['method_used']}")
    print(f"   📊 Paragraphs: {extraction_stats['paragraphs_found']}, Tables: {extraction_stats['tables_found']}")
    
    return cleaned_content, extraction_stats

# === MULTIPLE CHUNKING STRATEGIES === #
class ChunkingStrategies:
    @staticmethod
    def word_based(text, chunk_size=300, overlap=50):
        """Word-based chunking with overlap"""
        words = text.split()
        chunks = []
        for i in range(0, len(words), chunk_size - overlap):
            chunk = ' '.join(words[i:i + chunk_size])
            if chunk.strip():
                chunks.append(chunk.strip())
        return chunks
    
    @staticmethod
    def sentence_based(text, sentences_per_chunk=5):
        """Sentence-based chunking"""
        # Simple sentence splitting
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        chunks = []
        for i in range(0, len(sentences), sentences_per_chunk):
            chunk = '. '.join(sentences[i:i + sentences_per_chunk])
            if chunk.strip():
                chunks.append(chunk.strip() + '.')
        return chunks
    
    @staticmethod
    def paragraph_based(text):
        """Paragraph-based chunking"""
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        return paragraphs
    
    @staticmethod
    def token_based(text, max_tokens=512, overlap_tokens=50, model="gpt-4"):
        """Token-based chunking"""
        try:
            if model.startswith("gpt-4"):
                encoding = tiktoken.encoding_for_model("gpt-4")
            else:
                encoding = tiktoken.get_encoding("cl100k_base")
            
            tokens = encoding.encode(text)
            chunks = []
            
            for i in range(0, len(tokens), max_tokens - overlap_tokens):
                chunk_tokens = tokens[i:i + max_tokens]
                chunk_text = encoding.decode(chunk_tokens)
                chunks.append(chunk_text)
            
            return chunks
        except:
            # Fallback to word-based
            return ChunkingStrategies.word_based(text, chunk_size=int(max_tokens * 0.75))

# === COMPREHENSIVE RAG ANALYZER === #
class RAGStrategyAnalyzer:
    def __init__(self, model_name="gpt-4"):
        self.model_name = model_name
        self.max_context_tokens = MODEL_LIMITS.get(model_name, 8192)
        self.embedder = SentenceTransformer("all-MiniLM-L6-v2")
        self.documents = {}
        self.strategies = {}
        self.embeddings = {}
        self.analysis_results = {}
        
    def load_documents(self, filepaths):
        """Load multiple documents with detailed analysis"""
        print(f"\n🔄 Loading documents with {self.model_name} optimization...")
        
        for name, filepath in filepaths.items():
            content, stats = extract_text_robust(filepath)
            if content:
                self.documents[name] = {
                    'content': content,
                    'stats': stats,
                    'word_count': len(content.split()),
                    'token_count': count_tokens(content, self.model_name)
                }
                print(f"   📄 {name}: {stats['raw_length']} chars → {stats['cleaned_length']} chars")
    
    def apply_chunking_strategies(self):
        """Apply multiple chunking strategies to all documents"""
        print("\n🧠 Applying chunking strategies...")
        
        for doc_name, doc_data in self.documents.items():
            content = doc_data['content']
            self.strategies[doc_name] = {}
            
            # Apply different strategies
            strategies = {
                'word_300': ChunkingStrategies.word_based(content, 300, 50),
                'word_500': ChunkingStrategies.word_based(content, 500, 75),
                'sentence_5': ChunkingStrategies.sentence_based(content, 5),
                'sentence_8': ChunkingStrategies.sentence_based(content, 8),
                'paragraph': ChunkingStrategies.paragraph_based(content),
                'token_512': ChunkingStrategies.token_based(content, 512, 50, self.model_name),
                'token_256': ChunkingStrategies.token_based(content, 256, 25, self.model_name)
            }
            
            for strategy_name, chunks in strategies.items():
                token_analysis = analyze_token_distribution(chunks, self.model_name)
                self.strategies[doc_name][strategy_name] = {
                    'chunks': chunks,
                    'chunk_count': len(chunks),
                    'token_analysis': token_analysis
                }
            
            print(f"   📋 {doc_name}: {len(strategies)} strategies applied")
    
    def generate_embeddings(self):
        """Generate embeddings for all strategies"""
        print("\n🤖 Generating embeddings...")
        
        for doc_name, doc_strategies in self.strategies.items():
            self.embeddings[doc_name] = {}
            for strategy_name, strategy_data in doc_strategies.items():
                chunks = strategy_data['chunks']
                if chunks:
                    embeddings = self.embedder.encode(chunks)
                    self.embeddings[doc_name][strategy_name] = embeddings
                    print(f"   ✅ {doc_name}/{strategy_name}: {len(chunks)} chunks embedded")
    
    def analyze_strategy_performance(self, test_questions):
        """Analyze performance of different strategies"""
        print(f"\n🔍 Analyzing strategy performance with {len(test_questions)} test questions...")
        
        results = defaultdict(lambda: defaultdict(list))
        
        for question in test_questions:
            q_vec = self.embedder.encode([question])
            
            for doc_name, doc_embeddings in self.embeddings.items():
                for strategy_name, embeddings in doc_embeddings.items():
                    if len(embeddings) > 0:
                        similarities = cosine_similarity(q_vec, embeddings)[0]
                        top_indices = similarities.argsort()[-5:][::-1]
                        
                        # Get top chunks
                        chunks = self.strategies[doc_name][strategy_name]['chunks']
                        top_chunks = [chunks[i] for i in top_indices]
                        top_similarities = [similarities[i] for i in top_indices]
                        
                        # Calculate context size
                        context = '\n'.join(top_chunks)
                        context_tokens = count_tokens(context, self.model_name)
                        
                        # Calculate efficiency metrics
                        avg_similarity = np.mean(top_similarities)
                        efficiency = avg_similarity / context_tokens if context_tokens > 0 else 0
                        
                        results[doc_name][strategy_name].append({
                            'question': question,
                            'context_tokens': context_tokens,
                            'avg_similarity': avg_similarity,
                            'efficiency': efficiency,
                            'chunks_used': len(top_chunks)
                        })
        
        return results
    
    def get_strategy_recommendations(self, performance_results):
        """Get recommendations for best strategies"""
        print("\n📊 Calculating strategy recommendations...")
        
        recommendations = {}
        
        for doc_name, doc_results in performance_results.items():
            strategy_scores = {}
            
            for strategy_name, question_results in doc_results.items():
                avg_efficiency = np.mean([r['efficiency'] for r in question_results])
                avg_tokens = np.mean([r['context_tokens'] for r in question_results])
                avg_similarity = np.mean([r['avg_similarity'] for r in question_results])
                
                # Calculate composite score
                token_utilization = min(avg_tokens / self.max_context_tokens, 1.0)
                composite_score = avg_similarity * 0.5 + avg_efficiency * 0.3 + (1 - token_utilization) * 0.2
                
                strategy_scores[strategy_name] = {
                    'composite_score': composite_score,
                    'avg_efficiency': avg_efficiency,
                    'avg_tokens': avg_tokens,
                    'avg_similarity': avg_similarity,
                    'token_utilization': token_utilization
                }
            
            # Sort by composite score
            sorted_strategies = sorted(strategy_scores.items(), key=lambda x: x[1]['composite_score'], reverse=True)
            recommendations[doc_name] = sorted_strategies
        
        return recommendations
    
    def display_comprehensive_analysis(self, performance_results, recommendations):
        """Display comprehensive analysis results"""
        print("\n" + "="*80)
        print("🎯 COMPREHENSIVE RAG STRATEGY ANALYSIS")
        print("="*80)
        
        print(f"📋 Model: {self.model_name}")
        print(f"🔢 Max Context Tokens: {self.max_context_tokens:,}")
        print(f"📚 Documents Analyzed: {len(self.documents)}")
        
        # Document overview
        print(f"\n📄 DOCUMENT OVERVIEW:")
        for doc_name, doc_data in self.documents.items():
            print(f"   {doc_name}:")
            print(f"     Words: {doc_data['word_count']:,}")
            print(f"     Tokens: {doc_data['token_count']:,}")
            print(f"     Max Possible Context: {min(doc_data['token_count'], self.max_context_tokens):,} tokens")
        
        # Strategy performance
        print(f"\n🏆 STRATEGY RANKINGS:")
        for doc_name, strategy_ranking in recommendations.items():
            print(f"\n   📋 {doc_name.upper()}:")
            for i, (strategy_name, metrics) in enumerate(strategy_ranking[:3], 1):
                print(f"     {i}. {strategy_name}")
                print(f"        Score: {metrics['composite_score']:.4f}")
                print(f"        Avg Similarity: {metrics['avg_similarity']:.4f}")
                print(f"        Avg Tokens: {metrics['avg_tokens']:.0f}")
                print(f"        Token Utilization: {metrics['token_utilization']:.1%}")
                print(f"        Efficiency: {metrics['avg_efficiency']:.6f}")
        
        # Token usage analysis
        print(f"\n📊 TOKEN USAGE ANALYSIS:")
        for doc_name, doc_strategies in self.strategies.items():
            print(f"\n   📋 {doc_name}:")
            for strategy_name, strategy_data in doc_strategies.items():
                ta = strategy_data['token_analysis']
                print(f"     {strategy_name}:")
                print(f"       Chunks: {strategy_data['chunk_count']}")
                print(f"       Total Tokens: {ta['total_tokens']:,}")
                print(f"       Avg/Chunk: {ta['avg_tokens_per_chunk']:.1f}")
                print(f"       Max/Chunk: {ta['max_tokens_per_chunk']}")
                print(f"       Potential Context: {min(ta['total_tokens'], self.max_context_tokens):,}")
        
        # Optimization recommendations
        print(f"\n💡 OPTIMIZATION RECOMMENDATIONS:")
        for doc_name, strategy_ranking in recommendations.items():
            best_strategy = strategy_ranking[0]
            strategy_name, metrics = best_strategy
            
            print(f"\n   📋 {doc_name}:")
            print(f"     🏆 Best Strategy: {strategy_name}")
            print(f"     📊 Expected Context Size: {metrics['avg_tokens']:.0f} tokens")
            print(f"     🎯 Expected Relevance: {metrics['avg_similarity']:.1%}")
            print(f"     ⚡ Efficiency: {metrics['avg_efficiency']:.6f}")
            
            # Usage percentage
            usage_pct = (metrics['avg_tokens'] / self.max_context_tokens) * 100
            if usage_pct > 80:
                print(f"     ⚠️  High token usage ({usage_pct:.1f}% of limit)")
            elif usage_pct < 20:
                print(f"     ✅ Low token usage ({usage_pct:.1f}% of limit) - room for more context")
            else:
                print(f"     ✅ Optimal token usage ({usage_pct:.1f}% of limit)")

# === MAIN EXECUTION === #
def main():
    # Initialize analyzer
    analyzer = RAGStrategyAnalyzer(MODEL_NAME)
    
    # Document paths
    filepaths = {
        'cx_docs': 'CX_Responses.docx',
        'historical': 'Hoichoi_Export_497_tickets_20250718_155201.docx'
    }
    
    # Load documents
    analyzer.load_documents(filepaths)
    
    if not analyzer.documents:
        print("❌ No documents loaded successfully.")
        return
    
    # Apply chunking strategies
    analyzer.apply_chunking_strategies()
    
    # Generate embeddings
    analyzer.generate_embeddings()
    
    # Test questions for analysis
    test_questions = [
        "How do I reset my password?",
        "Video not playing properly",
        "Payment failed error",
        "Account login issues",
        "Streaming quality problems",
        "Subscription not working",
        "App crashes frequently",
        "Content not loading"
    ]
    
    # Analyze performance
    performance_results = analyzer.analyze_strategy_performance(test_questions)
    
    # Get recommendations
    recommendations = analyzer.get_strategy_recommendations(performance_results)
    
    # Display comprehensive analysis
    analyzer.display_comprehensive_analysis(performance_results, recommendations)
    
    print(f"\n🎉 Analysis complete! Best strategies identified for optimal token usage.")

if __name__ == "__main__":
    main()